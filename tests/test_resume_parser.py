from io import BytesIO
from pathlib import Path

from app.intake import resume_parser as rp


def test_extract_text_plaintext():
    data = "王小明\nPython 後端工程師".encode("utf-8")
    assert "Python" in rp.extract_text(data, "resume.txt")


def test_extract_text_unknown_ext_treated_as_text():
    data = "純文字履歷".encode("utf-8")
    assert rp.extract_text(data, "resume.unknown") == "純文字履歷"


def test_extract_text_docx():
    from docx import Document
    doc = Document()
    doc.add_paragraph("王小明")
    doc.add_paragraph("Python 後端工程師，3 年經驗")
    buf = BytesIO()
    doc.save(buf)
    text = rp.extract_text(buf.getvalue(), "resume.docx")
    assert "王小明" in text
    assert "後端工程師" in text


def test_extract_text_pdf():
    data = Path("tests/fixtures/sample_resume.pdf").read_bytes()
    text = rp.extract_text(data, "resume.pdf")
    assert "Resume" in text


def test_extract_text_strips_null_bytes():
    # UTF-16 存的 .txt、或某些 PDF 抽取結果會殘留 NUL(\x00)。\x00 是合法 UTF-8(U+0000)，
    # errors="ignore" 不會清掉它；接著被當 CLI 參數丟進 subprocess 會直接 ValueError("embedded null byte")。
    text = rp.extract_text(b"A\x00B\x00C", "resume.txt")
    assert "\x00" not in text
    assert "ABC" in text


def test_extract_text_preserves_whitespace():
    # 清控制字元時不可誤殺常見空白（換行/Tab），否則履歷排版全毀。
    text = rp.extract_text("第一行\n第二行\t縮排".encode("utf-8"), "resume.txt")
    assert "\n" in text and "\t" in text
